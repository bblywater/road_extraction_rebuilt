from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--report", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    project_root = Path(__file__).resolve().parent.parent
    report_path = (project_root / args.report).resolve()
    output_path = (project_root / args.output).resolve()

    document = Document()
    for line in report_path.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        else:
            document.add_paragraph(line)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print({"output": str(output_path)})


if __name__ == "__main__":
    main()
